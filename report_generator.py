import os
from pathlib import Path
from typing import List, Dict, Any, Optional

# Document generation packages
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

from config import REPORT_DIR
from logger import logger
from models import DDRReport

# ==========================================
# Helpers for Image Resolution
# ==========================================
def find_image_path(image_id: str, images_metadata: List[Dict[str, Any]]) -> Optional[str]:
    """
    Looks up an image ID in the extracted image metadata to find its file path.
    
    Args:
        image_id (str): The ID of the image to find.
        images_metadata (List[Dict[str, Any]]): The list of image metadata dictionaries.
        
    Returns:
        Optional[str]: The absolute path to the image file, or None if not found/invalid.
    """
    if not image_id:
        return None
    for img in images_metadata:
        if img.get("image_id") == image_id:
            path_str = img.get("path")
            if path_str and Path(path_str).exists():
                return path_str
    return None

# ==========================================
# DOCX Generation
# ==========================================
def generate_docx_report(ddr: DDRReport, images_metadata: List[Dict[str, Any]], filename: str = "DDR_Report.docx") -> Path:
    """
    Generates a professionally styled DOCX report based on DDRReport schema.
    
    Args:
        ddr (DDRReport): The validated DDR Report data.
        images_metadata (List[Dict[str, Any]]): Image metadata for resolving file paths.
        filename (str): Name of output file.
        
    Returns:
        Path: Path to the generated DOCX report.
    """
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = REPORT_DIR / filename
    logger.info(f"Generating DOCX report at: {docx_path}")
    
    doc = Document()
    
    # Define primary theme colors (Navy and Dark Grey)
    PRIMARY_COLOR = RGBColor(11, 29, 58)  # Deep Navy
    SECONDARY_COLOR = RGBColor(108, 117, 125)  # Grey
    TEXT_COLOR = RGBColor(33, 37, 41)  # Charcoal
    
    # Configure base styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = TEXT_COLOR
    
    # ----------------------------------------------------
    # 1. Cover Page
    # ----------------------------------------------------
    # Header spacing
    for _ in range(3):
        doc.add_paragraph()
        
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("DETAILED DIAGNOSTIC REPORT")
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Professional Building Diagnostics & Thermal Analysis")
    sub_run.font.size = Pt(14)
    sub_run.italic = True
    sub_run.font.color.rgb = SECONDARY_COLOR
    
    # Decorative line
    border_p = doc.add_paragraph()
    border_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_run = border_p.add_run("____________________________________________________")
    border_run.font.color.rgb = PRIMARY_COLOR
    
    for _ in range(5):
        doc.add_paragraph()
        
    # Property Info Metadata Table
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_table.style = 'Table Grid'
    
    labels = ["Property Address:", "Client/Owner:", "Inspection Date:"]
    values = [ddr.property_info.address, ddr.property_info.client_name, ddr.property_info.inspection_date]
    
    for idx, (label, val) in enumerate(zip(labels, values)):
        row = meta_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(str(val))
        
    doc.add_page_break()
    
    # ----------------------------------------------------
    # 2. Executive Summary
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Executive Summary")
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1_run.bold = True
    
    doc.add_paragraph(ddr.executive_summary)
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # 3. Overall Severity Assessment & Building Health
    # ----------------------------------------------------
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Property Severity & Building Health Assessment")
    h2_run.font.color.rgb = PRIMARY_COLOR
    h2_run.bold = True
    
    severity_p = doc.add_paragraph()
    severity_p.add_run("Overall Property Severity: ").bold = True
    severity_run = severity_p.add_run(ddr.severity_assessment.overall_severity.upper())
    severity_run.bold = True
    if ddr.severity_assessment.overall_severity.lower() in ['critical', 'high']:
        severity_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
    else:
        severity_run.font.color.rgb = RGBColor(40, 167, 69)  # Green
        
    risk_p = doc.add_paragraph()
    risk_p.add_run("Calculated Risk Score (0-100): ").bold = True
    risk_p.add_run(f"{ddr.severity_assessment.risk_score}/100")
    
    doc.add_paragraph(ddr.severity_assessment.reasoning)
    
    # Health Scores Table
    doc.add_paragraph("Building Health Indices:").bold = True
    health_table = doc.add_table(rows=5, cols=2)
    health_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    health_table.style = 'Table Grid'
    
    h_labels = ["Overall Building Health Score", "Structural Health Index", "Waterproofing Health Index", "Interior Finish Health Index", "Overall Moisture Risk Level"]
    h_values = [f"{100 - ddr.severity_assessment.risk_score}%", f"{ddr.severity_assessment.structural_health}%", f"{ddr.severity_assessment.waterproofing_health}%", f"{ddr.severity_assessment.interior_finish_health}%", ddr.severity_assessment.moisture_risk]
    
    for idx, (label, val) in enumerate(zip(h_labels, h_values)):
        row = health_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(str(val))
        
    doc.add_paragraph()
    
    # Priority matrix / ranking
    doc.add_paragraph("Area Repair Priority Ranking (Highest to Lowest):").bold = True
    for rank, area in enumerate(ddr.severity_assessment.priority_ranking):
        doc.add_paragraph(f"{rank + 1}. {area}", style='List Bullet')
        
    # Cost ranges
    cost_p = doc.add_paragraph()
    cost_p.add_run("Total Estimated Remediation Program Cost: ").bold = True
    cost_p.add_run(ddr.estimated_total_cost_range).bold = True
    
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # 4. Area-wise Observations
    # ----------------------------------------------------
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Area-wise Observations & Diagnostics")
    h3_run.font.color.rgb = PRIMARY_COLOR
    h3_run.bold = True
    
    for obs in ddr.observations:
        # Obs Heading
        obs_h = doc.add_heading(level=2)
        obs_run = obs_h.add_run(f"3.{obs.id} {obs.area} - {obs.defect}")
        obs_run.font.color.rgb = PRIMARY_COLOR
        
        # Details
        doc.add_paragraph(obs.description)
        
        # Metadata details
        p_meta = doc.add_paragraph()
        p_meta.add_run("Severity: ").bold = True
        sev_r = p_meta.add_run(obs.severity)
        sev_r.bold = True
        
        p_meta.add_run("   |   Confidence: ").bold = True
        conf_r = p_meta.add_run(obs.confidence)
        conf_r.bold = True
        
        p_meta.add_run("   |   Probable Root Cause: ").bold = True
        p_meta.add_run(obs.probable_root_cause)
        
        # Priority and Timeline
        p_meta2 = doc.add_paragraph()
        p_meta2.add_run("Priority Level: ").bold = True
        p_meta2.add_run(obs.priority)
        p_meta2.add_run("   |   Timeline: ").bold = True
        p_meta2.add_run(obs.timeline)
        p_meta2.add_run("   |   Estimated Repair Cost: ").bold = True
        p_meta2.add_run(obs.estimated_cost)
        
        # Conflict resolution if any
        if obs.is_conflict:
            conf_p = doc.add_paragraph()
            conf_p.add_run("⚠ Conflict Detected: ").bold = True
            conf_p.add_run(str(obs.conflict_details)).italic = True
            
        # Sequential Repair Procedure
        doc.add_paragraph("Sequential Engineering Repair Procedure:").bold = True
        for step in obs.repair_procedure:
            doc.add_paragraph(step, style='List Bullet')
            
        # Materials Required
        p_mats = doc.add_paragraph()
        p_mats.add_run("Materials Required: ").bold = True
        p_mats.add_run(", ".join(obs.materials_required))
        
        # Confidence Breakdown
        p_conf = doc.add_paragraph()
        p_conf.add_run("AI Confidence Telemetry: ").bold = True
        c_parts = [f"{k}: {v}%" for k, v in obs.confidence_breakdown.items()]
        p_conf.add_run("   |   ".join(c_parts))
        
        # Evidence sources
        doc.add_paragraph("Evidence Traceability:").bold = True
        for ev in obs.evidence:
            ev_str = f"- {ev.source_document} (Page {ev.page_number})"
            doc.add_paragraph(ev_str, style='List Bullet')
            
            # Embed image if ID matches
            if ev.image_id:
                for img_code in ev.image_id.split(','):
                    img_code_clean = img_code.strip()
                    if img_code_clean:
                        img_path = find_image_path(img_code_clean, images_metadata)
                        if img_path:
                            try:
                                doc.add_picture(img_path, width=Inches(4.5))
                                caption_p = doc.add_paragraph()
                                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cap_r = caption_p.add_run(f"Figure: {img_code_clean} - Associated image extracted from {ev.source_document} (Page {ev.page_number})")
                                cap_r.italic = True
                                cap_r.font.size = Pt(9.5)
                            except Exception as err:
                                logger.error(f"Failed to insert image {img_code_clean} in DOCX: {str(err)}")
                                doc.add_paragraph(f"[Image {img_code_clean} Not Available]")
                        
        doc.add_paragraph()
        
    # ----------------------------------------------------
    # 5. Final Engineering Assessment
    # ----------------------------------------------------
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Final Engineering Assessment")
    h4_run.font.color.rgb = PRIMARY_COLOR
    h4_run.bold = True
    
    doc.add_paragraph(ddr.final_engineering_assessment)
    doc.add_paragraph()

    # ----------------------------------------------------
    # 6. Recommended Actions
    # ----------------------------------------------------
    h5 = doc.add_heading(level=1)
    h5_run = h5.add_run("5. Recommended Corrective Actions")
    h5_run.font.color.rgb = PRIMARY_COLOR
    h5_run.bold = True
    
    # Group recommendations by priority
    recs_by_priority: Dict[str, List[Any]] = {
        "Immediate Actions": [],
        "Short-term Actions": [],
        "Long-term Preventive Actions": []
    }
    
    for rec in ddr.recommendations:
        priority = rec.priority
        if priority in recs_by_priority:
            recs_by_priority[priority].append(rec)
        else:
            # Fallback mapping
            recs_by_priority["Long-term Preventive Actions"].append(rec)
            
    for priority, rec_list in recs_by_priority.items():
        doc.add_heading(priority, level=2)
        if not rec_list:
            doc.add_paragraph("No actions required in this category.")
        else:
            for rec in rec_list:
                associated = ", ".join(rec.associated_observation_ids)
                p_rec = doc.add_paragraph(style='List Bullet')
                p_rec.add_run(f"[{rec.action_type}] ").bold = True
                p_rec.add_run(f"{rec.description} (Addresses: {associated})")
                
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # 7. Missing or Unclear Information
    # ----------------------------------------------------
    h6 = doc.add_heading(level=1)
    h6_run = h6.add_run("6. Missing or Unclear Information")
    h6_run.font.color.rgb = PRIMARY_COLOR
    h6_run.bold = True
    
    if not ddr.missing_information:
        doc.add_paragraph("No significant missing information identified. The documentation provided was sufficient.")
    else:
        for item in ddr.missing_information:
            p_miss = doc.add_paragraph(style='List Bullet')
            p_miss.add_run(f"{item.item_category}: ").bold = True
            p_miss.add_run(item.description)
            
    try:
        doc.save(docx_path)
        logger.info(f"DOCX report saved successfully.")
        return docx_path
    except Exception as e:
        logger.error(f"Failed to save DOCX file: {str(e)}")
        raise e

# ==========================================
# PDF Generation
# ==========================================
def draw_page_number(canvas, doc):
    """
    Draws footers with page numbers on all PDF pages.
    """
    canvas.saveState()
    canvas.setFont('Helvetica', 9)
    canvas.setFillColor(colors.HexColor("#6c757d"))
    # Right-aligned page numbers
    page_num = canvas.getPageNumber()
    canvas.drawRightString(8.0 * inch, 0.5 * inch, f"Page {page_num}")
    # Left-aligned header
    canvas.drawString(0.5 * inch, 10.5 * inch, "Detailed Diagnostic Report (DDR) - Building Diagnostics")
    canvas.setStrokeColor(colors.HexColor("#dee2e6"))
    canvas.setLineWidth(0.5)
    canvas.line(0.5 * inch, 10.4 * inch, 8.0 * inch, 10.4 * inch)
    canvas.restoreState()

def generate_pdf_report(ddr: DDRReport, images_metadata: List[Dict[str, Any]], filename: str = "DDR_Report.pdf") -> Path:
    """
    Generates a professionally styled PDF report using ReportLab.
    
    Args:
        ddr (DDRReport): The validated DDR Report data.
        images_metadata (List[Dict[str, Any]]): Image metadata for resolving file paths.
        filename (str): Name of output file.
        
    Returns:
        Path: Path to the generated PDF report.
    """
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    pdf_path = REPORT_DIR / filename
    logger.info(f"Generating PDF report at: {pdf_path}")
    
    # Setup document document template with 0.5 inch margins
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=26,
        leading=32,
        textColor=colors.HexColor('#0b1d3a'),
        alignment=1, # Centered
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#6c757d'),
        alignment=1, # Centered
        spaceAfter=40
    )
    
    h1_style = ParagraphStyle(
        'SecHeading',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#0b1d3a'),
        spaceBefore=15,
        spaceAfter=10,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'SubSecHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#2c3e50'),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#212529'),
        spaceAfter=8
    )
    
    bold_body_style = ParagraphStyle(
        'BoldBody',
        parent=body_style,
        fontName='Helvetica-Bold'
    )
    
    caption_style = ParagraphStyle(
        'Caption',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#495057'),
        alignment=1,
        spaceAfter=12
    )

    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=12
    )
    
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=12,
        textColor=colors.white
    )

    story = []
    
    # ----------------------------------------------------
    # 1. Cover Page
    # ----------------------------------------------------
    story.append(Spacer(1, 1.5 * inch))
    story.append(Paragraph("DETAILED DIAGNOSTIC REPORT", title_style))
    story.append(Paragraph("Professional Building Diagnostics & Thermal Inspection Summary", subtitle_style))
    story.append(Spacer(1, 1.0 * inch))
    
    # Meta Info block
    meta_data = [
        [Paragraph("Property Address:", bold_body_style), Paragraph(ddr.property_info.address, body_style)],
        [Paragraph("Client Name:", bold_body_style), Paragraph(ddr.property_info.client_name, body_style)],
        [Paragraph("Inspection Date:", bold_body_style), Paragraph(ddr.property_info.inspection_date, body_style)]
    ]
    meta_table = Table(meta_data, colWidths=[2.0 * inch, 4.5 * inch])
    meta_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#dee2e6")),
        ('PADDING', (0,0), (-1,-1), 8),
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#f8f9fa")),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    
    story.append(meta_table)
    story.append(PageBreak())
    
    # ----------------------------------------------------
    # 2. Executive Summary
    # ----------------------------------------------------
    story.append(Paragraph("1. Executive Summary", h1_style))
    story.append(Paragraph(ddr.executive_summary, body_style))
    story.append(Spacer(1, 10))
    
    # ----------------------------------------------------
    # 3. Severity & Property Risk Assessment
    # ----------------------------------------------------
    story.append(Paragraph("2. Severity & Property Risk Assessment", h1_style))
    
    # Text metrics
    story.append(Paragraph(f"<b>Overall Property Severity:</b> {ddr.severity_assessment.overall_severity.upper()}", body_style))
    story.append(Paragraph(f"<b>Calculated Risk Score:</b> {ddr.severity_assessment.risk_score} / 100", body_style))
    story.append(Paragraph(ddr.severity_assessment.reasoning, body_style))
    story.append(Spacer(1, 5))
    
    # Health indices table
    story.append(Paragraph("<b>Building Health Indices:</b>", body_style))
    health_grid = [
        [Paragraph("<b>Metric</b>", table_header_style), Paragraph("<b>Value</b>", table_header_style)],
        [Paragraph("Overall Building Health Score", table_cell_style), Paragraph(f"{100 - ddr.severity_assessment.risk_score}%", table_cell_style)],
        [Paragraph("Structural Health Index", table_cell_style), Paragraph(f"{ddr.severity_assessment.structural_health}%", table_cell_style)],
        [Paragraph("Waterproofing Health Index", table_cell_style), Paragraph(f"{ddr.severity_assessment.waterproofing_health}%", table_cell_style)],
        [Paragraph("Interior Finish Health Index", table_cell_style), Paragraph(f"{ddr.severity_assessment.interior_finish_health}%", table_cell_style)],
        [Paragraph("Moisture Risk Level", table_cell_style), Paragraph(ddr.severity_assessment.moisture_risk, table_cell_style)],
        [Paragraph("Estimated Total Cost Range", table_cell_style), Paragraph(ddr.estimated_total_cost_range, table_cell_style)]
    ]
    h_table = Table(health_grid, colWidths=[3.2 * inch, 3.2 * inch])
    h_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#dee2e6")),
        ('PADDING', (0,0), (-1,-1), 6),
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#0b1d3a")),
        ('BACKGROUND', (0,1), (0,-1), colors.HexColor("#f8f9fa")),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(h_table)
    story.append(Spacer(1, 10))
    
    # Priority ranking bullet list
    story.append(Paragraph("<b>Area Action Priority Ranking:</b>", body_style))
    for idx, area in enumerate(ddr.severity_assessment.priority_ranking):
        story.append(Paragraph(f"&bull; Priority {idx + 1}: {area}", body_style))
    story.append(Spacer(1, 15))
    
    # ----------------------------------------------------
    # 4. Area-wise Observations
    # ----------------------------------------------------
    story.append(Paragraph("3. Detailed Area-wise Observations", h1_style))
    
    for obs in ddr.observations:
        obs_elements = []
        obs_elements.append(Paragraph(f"3.{obs.id} {obs.area} - {obs.defect}", h2_style))
        obs_elements.append(Paragraph(obs.description, body_style))
        
        # Details grid table
        meta_grid = [
            [Paragraph("<b>Severity:</b>", table_cell_style), Paragraph(obs.severity, table_cell_style),
             Paragraph("<b>Confidence:</b>", table_cell_style), Paragraph(obs.confidence, table_cell_style)],
            [Paragraph("<b>Priority:</b>", table_cell_style), Paragraph(obs.priority, table_cell_style),
             Paragraph("<b>Timeline:</b>", table_cell_style), Paragraph(obs.timeline, table_cell_style)],
            [Paragraph("<b>Root Cause:</b>", table_cell_style), Paragraph(obs.probable_root_cause, table_cell_style),
             Paragraph("<b>Estimated Cost:</b>", table_cell_style), Paragraph(obs.estimated_cost, table_cell_style)],
            [Paragraph("<b>Conflict Status:</b>", table_cell_style), Paragraph("Yes" if obs.is_conflict else "No", table_cell_style),
             Paragraph("<b>Materials:</b>", table_cell_style), Paragraph(", ".join(obs.materials_required), table_cell_style)]
        ]
        grid_table = Table(meta_grid, colWidths=[1.2 * inch, 2.0 * inch, 1.2 * inch, 2.1 * inch])
        grid_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e9ecef")),
            ('PADDING', (0,0), (-1,-1), 5),
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#f8f9fa")),
            ('BACKGROUND', (2,0), (2,-1), colors.HexColor("#f8f9fa")),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        
        obs_elements.append(grid_table)
        obs_elements.append(Spacer(1, 8))
        
        if obs.is_conflict:
            obs_elements.append(Paragraph(f"<b>Conflict Note:</b> {obs.conflict_details}", body_style))
            obs_elements.append(Spacer(1, 5))
            
        # Repair Procedure
        obs_elements.append(Paragraph("<b>Engineering Repair Procedure:</b>", body_style))
        for step in obs.repair_procedure:
            obs_elements.append(Paragraph(f"&bull; {step}", body_style))
        obs_elements.append(Spacer(1, 5))
        
        # Confidence breakdown
        c_parts = [f"{k}: {v}%" for k, v in obs.confidence_breakdown.items()]
        obs_elements.append(Paragraph(f"<b>AI Confidence Breakdown:</b> { ' | '.join(c_parts) }", body_style))
        obs_elements.append(Spacer(1, 5))

        # Evidence and pictures
        obs_elements.append(Paragraph("<b>Evidence Sources:</b>", body_style))
        for ev in obs.evidence:
            obs_elements.append(Paragraph(f"&bull; {ev.source_document} (Page {ev.page_number})", body_style))
            
            # Add Image(s)
            if ev.image_id:
                for img_code in ev.image_id.split(','):
                    img_code_clean = img_code.strip()
                    if img_code_clean:
                        img_path = find_image_path(img_code_clean, images_metadata)
                        if img_path:
                            try:
                                # Add Image flowable, scaled nicely to fit width
                                img_flow = Image(img_path, width=4.0 * inch, height=3.0 * inch)
                                img_flow.hAlign = 'CENTER'
                                obs_elements.append(Spacer(1, 5))
                                obs_elements.append(img_flow)
                                obs_elements.append(Paragraph(f"Figure: {img_code_clean} - Associated image extracted from {ev.source_document} (Page {ev.page_number})", caption_style))
                            except Exception as img_err:
                                logger.error(f"Failed to embed image {img_code_clean} in PDF: {str(img_err)}")
                                obs_elements.append(Paragraph(f"<i>[Image {img_code_clean} Not Available]</i>", caption_style))
                        
        obs_elements.append(Spacer(1, 10))
        story.append(KeepTogether(obs_elements))
        
    # ----------------------------------------------------
    # 5. Final Engineering Assessment
    # ----------------------------------------------------
    story.append(PageBreak())
    story.append(Paragraph("4. Final Engineering Assessment", h1_style))
    story.append(Paragraph(ddr.final_engineering_assessment, body_style))
    story.append(Spacer(1, 15))
        
    # ----------------------------------------------------
    # 6. Recommendations
    # ----------------------------------------------------
    story.append(Paragraph("5. Recommended Corrective Actions", h1_style))
    
    # Priority subheadings
    rec_priorities = ["Immediate Actions", "Short-term Actions", "Long-term Preventive Actions"]
    for priority in rec_priorities:
        story.append(Paragraph(priority, h2_style))
        priority_recs = [r for r in ddr.recommendations if r.priority == priority]
        
        if not priority_recs:
            story.append(Paragraph("No corrective actions identified for this timeframe.", body_style))
        else:
            for rec in priority_recs:
                associated = ", ".join(rec.associated_observation_ids)
                rec_text = f"<b>[{rec.action_type}]</b> {rec.description} <i>(Resolves: {associated})</i>"
                story.append(Paragraph(f"&bull; {rec_text}", body_style))
                
        story.append(Spacer(1, 8))
        
    # ----------------------------------------------------
    # 7. Missing Information
    # ----------------------------------------------------
    story.append(Paragraph("6. Missing or Unclear Information", h1_style))
    
    if not ddr.missing_information:
        story.append(Paragraph("All documents and measurements necessary for compilation were fully available.", body_style))
    else:
        for item in ddr.missing_information:
            story.append(Paragraph(f"&bull; <b>{item.item_category}:</b> {item.description}", body_style))
            
    # Build the document
    try:
        # Build using our custom canvas callback for running page numbers
        doc.build(story, onFirstPage=draw_page_number, onLaterPages=draw_page_number)
        logger.info(f"PDF report saved successfully.")
        return pdf_path
    except Exception as e:
        logger.error(f"Failed to build PDF document: {str(e)}")
        raise e
